# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# === Título principal ===
titulo = doc.add_heading('Cat Jard - Documentación de Base de Datos', level=0)
titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Webapp B2B de merchandising corporativo. Arquitectura de microservicios '
    'con Spring Boot + Spring Cloud y un frontend React. La base de datos sigue '
    'el patrón Database per Service sobre PostgreSQL.'
)

# ============================================================
# 2. INTEGRACIÓN BD
# ============================================================
doc.add_heading('2. Integración de Base de Datos', level=1)

doc.add_heading('2.1 Plataforma', level=2)
doc.add_paragraph(
    'Motor de base de datos: PostgreSQL 15+ (instancia local en puerto 5432). '
    'Se aplica el patrón Database per Service: cada microservicio tiene su propio '
    'esquema lógico independiente, lo que permite desplegar, versionar y escalar '
    'cada dominio por separado.'
)

p = doc.add_paragraph()
p.add_run('Bases de datos del sistema (definidas en ').bold = False
p.add_run('init-databases.sql').italic = True
p.add_run('):')

bases = [
    ('catjard_identity', 'Usuarios, roles, autenticación.'),
    ('catjard_catalog', 'Productos, categorías, técnicas, promociones.'),
    ('catjard_crm', 'Leads y clientes corporativos (CRM).'),
    ('catjard_sales', 'Cotizaciones y pedidos.'),
    ('catjard_inventory', 'Proveedores, órdenes de compra y movimientos de stock.'),
    ('catjard_operations', 'Artes gráficos, tracking de producción y despachos.'),
]
for nombre, desc in bases:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f' — {desc}')

doc.add_heading('2.2 Arquitectura de conexión', level=2)
doc.add_paragraph(
    'La integración se compone de tres capas: cliente web (React), capa de '
    'orquestación (API Gateway + Eureka) y capa de microservicios que persisten '
    'cada uno en su PostgreSQL.'
)

p = doc.add_paragraph()
p.add_run('Flujo de una petición típica: ').bold = True
doc.add_paragraph(
    'Frontend (React, http://localhost:5173) → API Gateway (Spring Cloud Gateway, '
    'puerto 8080) → Eureka Service Discovery (puerto 8761) → microservicio destino '
    '(identity 8081, catalog, crm, sales, inventory, operations) → PostgreSQL '
    '(puerto 5432, base correspondiente).'
)

doc.add_heading('Backend (Spring Boot)', level=3)
backend_items = [
    ('Spring Data JPA + Hibernate', 'capa ORM. ddl-auto=update en desarrollo, dialecto PostgreSQLDialect.'),
    ('Flyway', 'control de versiones del esquema (classpath:db/migration). baseline-on-migrate=true.'),
    ('HikariCP', 'pool de conexiones por defecto en Spring Boot.'),
    ('Spring Cloud Gateway', 'enruta /api/** hacia los servicios usando lb:// + Eureka.'),
    ('Eureka Client/Server', 'descubrimiento dinámico de instancias.'),
    ('Feign Client', 'comunicación inter-servicios (inventory → catalog, operations → sales) con interceptor que reenvía el JWT.'),
    ('Spring Security + JWT (HS256)', 'cada servicio valida el token con JwtAuthenticationFilter. El token incluye sub=email, userId, role, name.'),
    ('BCrypt', 'hash de contraseñas en identity-service.'),
    ('SpringDoc / Swagger UI', 'documentación automática en /swagger-ui.html por servicio.'),
]
for nombre, desc in backend_items:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

doc.add_heading('Frontend (React + Vite)', level=3)
frontend_items = [
    ('React 18 + Vite', 'SPA servida en localhost:5173.'),
    ('React Router 7', 'enrutamiento por roles: público, cliente, admin (ventas/almacén/operaciones/gerencia).'),
    ('Redux Toolkit', '15 slices de estado (auth, cart, productos, cotizaciones, pedidos, leads, clientes, etc.).'),
    ('Tailwind CSS', 'estilos utilitarios.'),
    ('apiClient.js', 'wrapper sobre fetch. Inyecta Authorization: Bearer <token> desde localStorage.'),
    ('ProtectedRoute', 'restringe rutas según el rol decodificado del JWT.'),
]
for nombre, desc in frontend_items:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

doc.add_heading('2.3 Organización de archivos del proyecto', level=2)

doc.add_paragraph('Estructura del backend (catjard/):', style='Intense Quote')
backend_tree = """catjard/
├── init-databases.sql          # crea las 6 BDs PostgreSQL
├── eureka-server/              # Service Discovery (puerto 8761)
├── api-gateway/                # Spring Cloud Gateway (puerto 8080)
│   └── src/main/resources/application.yml   # rutas /api/**, CORS
├── identity-service/           # Auth + usuarios (puerto 8081)
│   └── src/main/java/com/catjard/identity/
│       ├── controller/   AuthController, UsuarioController
│       ├── service/      AuthService, JwtService, UsuarioService, DataSeeder
│       ├── repository/   UsuarioRepository (Spring Data JPA)
│       ├── model/        Usuario, Rol (cliente|vendedor|almacen|produccion|gerente)
│       ├── dto/          LoginRequest, LoginResponse, RegistroDTO, UsuarioDTO...
│       ├── mapper/       UsuarioMapper
│       └── config/       SecurityConfig, JwtAuthenticationFilter, OpenApiConfig
├── catalog-service/            # Productos, categorías, técnicas, promociones
├── crm-service/                # Leads, ClienteCRM (razón social, RUC)
├── sales-service/              # Cotización, CotizacionItem, Pedido, PedidoItem
├── inventory-service/          # Proveedor, OrdenCompra, Movimiento
│   └── src/main/java/com/catjard/inventory/client/CatalogClient.java   # Feign
└── operations-service/         # Arte, TrackingEvento, Despacho
    └── src/main/java/com/catjard/operations/client/SalesClient.java    # Feign"""
p = doc.add_paragraph(backend_tree)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_paragraph('Estructura del frontend (datos/):', style='Intense Quote')
frontend_tree = """datos/
├── index.html
├── package.json                # React 18, Vite, Redux Toolkit, React Router 7
├── tailwind.config.js
├── vite.config.js
└── src/
    ├── App.jsx                 # Provider Redux + bootstrap fetchProductos
    ├── main.jsx
    ├── index.css
    ├── routes/
    │   ├── AppRouter.jsx       # rutas públicas / cliente / admin por rol
    │   └── ProtectedRoute.jsx  # gate por roles
    ├── layouts/
    │   ├── PublicLayout.jsx
    │   ├── ClientLayout.jsx
    │   └── AdminLayout.jsx
    ├── pages/
    │   ├── public/   Home, Catalogo, Login, Registro, RecuperarPassword...
    │   ├── client/   Dashboard, Carrito, Cotizaciones, Pedidos, Perfil
    │   └── admin/    ventas/, almacen/, operaciones/, gerencia/
    ├── components/   ProductCard, AdminHeader, Primitives, Icons...
    ├── services/     apiClient.js (fetch + JWT), authService.js,
    │                 productosService.js, cotizacionesService.js,
    │                 pedidosService.js, clientesService.js, leadsService.js,
    │                 artesService.js, trackingService.js, storage.js
    ├── redux/
    │   ├── store.js            # configureStore con 15 reducers
    │   └── slices/             # authSlice, cartSlice, productosSlice...
    ├── hooks/
    └── data/                   # mocks (clientes, productos, pedidos...)"""
p = doc.add_paragraph(frontend_tree)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_heading('2.4 Codificación: ejemplos relevantes', level=2)

doc.add_paragraph('Configuración de conexión (identity-service/application.properties):')
config_props = """spring.datasource.url=jdbc:postgresql://localhost:5432/catjard_identity
spring.datasource.username=postgres
spring.datasource.password=postgres
spring.datasource.driver-class-name=org.postgresql.Driver

spring.jpa.hibernate.ddl-auto=update
spring.jpa.properties.hibernate.dialect=org.hibernate.dialect.PostgreSQLDialect
spring.jpa.open-in-view=false

spring.flyway.enabled=true
spring.flyway.locations=classpath:db/migration
spring.flyway.baseline-on-migrate=true

jwt.secret=ZmI4NDdkN2I1NDhmNGIzN2JjMjNkMjJlOWE5MjUyNTI1MjMyNTNkMmFlMTNmNzJkZWE5MmU4MmM0YzY5OTRiNQ==
jwt.expiration-ms=86400000

eureka.client.service-url.defaultZone=http://localhost:8761/eureka/
eureka.instance.prefer-ip-address=true"""
p = doc.add_paragraph(config_props)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_paragraph('Entidad JPA (identity-service/model/Usuario.java):')
usuario_code = """@Entity
@Table(name = "usuarios")
@Getter @Setter @NoArgsConstructor @AllArgsConstructor @Builder
public class Usuario {
    @Id @GeneratedValue(strategy = GenerationType.IDENTITY)
    private Long id;

    @Column(nullable = false, unique = true)
    private String email;

    @Column(nullable = false)
    private String password;

    @Enumerated(EnumType.STRING)
    @Column(nullable = false, length = 20)
    private Rol rol;

    @Column(nullable = false)
    private String nombre;

    private String empresa;
    @Column(length = 11)  private String ruc;
    @Column(length = 30)  private String telefono;

    @Column(name = "fecha_creacion", nullable = false, updatable = false)
    private LocalDateTime fechaCreacion;
    @Column(name = "fecha_actualizacion", nullable = false)
    private LocalDateTime fechaActualizacion;

    @PrePersist void onCreate() { /* set timestamps */ }
    @PreUpdate  void onUpdate() { fechaActualizacion = LocalDateTime.now(); }
}"""
p = doc.add_paragraph(usuario_code)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_paragraph('Ruteo del API Gateway (api-gateway/application.yml, fragmento):')
gw_code = """spring:
  cloud:
    gateway:
      server:
        webflux:
          globalcors:
            cors-configurations:
              '[/**]':
                allowed-origins: ["http://localhost:5173"]
                allowed-methods: [GET, POST, PATCH, PUT, DELETE, OPTIONS]
                allow-credentials: true
          routes:
            - id: identity-auth
              uri: lb://identity-service
              predicates: [ Path=/api/auth/** ]
            - id: catalog-productos
              uri: lb://catalog-service
              predicates: [ Path=/api/productos/** ]
            - id: sales-cotizaciones
              uri: lb://sales-service
              predicates: [ Path=/api/cotizaciones/** ]"""
p = doc.add_paragraph(gw_code)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_paragraph('Cliente HTTP del frontend (src/services/apiClient.js):')
api_code = """const BASE_URL = import.meta.env.VITE_API_BASE_URL || 'http://localhost:8080/api';
const TOKEN_KEY = 'catjard_token';

async function request(method, path, { body, auth = true, headers = {} } = {}) {
  const finalHeaders = { ...headers };
  if (body !== undefined) finalHeaders['Content-Type'] = 'application/json';
  if (auth) {
    const token = getToken();
    if (token) finalHeaders['Authorization'] = `Bearer ${token}`;
  }
  const res = await fetch(`${BASE_URL}${path}`, {
    method, headers: finalHeaders,
    body: body !== undefined ? JSON.stringify(body) : undefined,
  });
  if (!res.ok) { /* manejo de error con status */ }
  return res.status === 204 ? null : res.json();
}

export const apiClient = {
  get:  (p, o) => request('GET',    p, o),
  post: (p, b, o) => request('POST',  p, { ...o, body: b }),
  put:  (p, b, o) => request('PUT',   p, { ...o, body: b }),
  del:  (p, o) => request('DELETE', p, o),
};"""
p = doc.add_paragraph(api_code)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_paragraph('Servicio de autenticación frontend (src/services/authService.js):')
auth_code = """export async function login({ email, password }) {
  const res = await apiClient.post('/auth/login', { email, password }, { auth: false });
  const token = res?.token || res?.accessToken || res?.jwt;
  if (!token) throw new Error('Respuesta del servidor sin token.');
  setToken(token);

  let user = res?.user;
  if (!user) {
    const claims = decodeJwt(token);
    user = {
      id:     claims.userId ?? claims.id,
      email:  claims.sub    ?? claims.email,
      nombre: claims.name   ?? claims.nombre,
      role:   claims.role   ?? claims.rol,
    };
  }
  return user;
}"""
p = doc.add_paragraph(auth_code)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

# ============================================================
# 3. REPLICACIÓN BD
# ============================================================
doc.add_heading('3. Replicación de Base de Datos', level=1)

doc.add_heading('3.1 Tipo de replicación aplicado', level=2)
doc.add_paragraph(
    'Se aplica replicación física asíncrona en streaming (PostgreSQL Streaming '
    'Replication) en topología primary-replica (un nodo primario que acepta '
    'lecturas y escrituras + uno o varios nodos réplica de solo lectura). '
    'La replicación se basa en el envío continuo de los registros WAL '
    '(Write-Ahead Log) desde el primario hacia las réplicas mediante el '
    'protocolo de replicación nativo de PostgreSQL.'
)

p = doc.add_paragraph()
p.add_run('Configuración resumida en el primario (postgresql.conf):').bold = True
prim = """wal_level = replica
max_wal_senders = 10
wal_keep_size = 1GB
archive_mode = on
archive_command = 'cp %p /var/lib/postgresql/wal_archive/%f'"""
p = doc.add_paragraph(prim)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

p = doc.add_paragraph()
p.add_run('En la réplica:').bold = True
rep = """primary_conninfo = 'host=primary_host port=5432 user=replicator password=...'
hot_standby = on"""
p = doc.add_paragraph(rep)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

doc.add_heading('3.2 Justificación', level=2)
just = [
    ('Alta disponibilidad (HA)',
     'si el nodo primario cae, una réplica puede ser promovida (failover) y el sistema '
     'sigue operativo. Para una webapp B2B que gestiona pedidos y cotizaciones de clientes '
     'corporativos, la indisponibilidad genera pérdida directa de ventas.'),
    ('Distribución de carga de lectura',
     'el frontend hace muchas más lecturas que escrituras (catálogo público, dashboards, '
     'consultas de pedidos). Las réplicas pueden absorber ese tráfico mientras el primario '
     'se enfoca en transacciones (cotizaciones, órdenes, movimientos de stock).'),
    ('Tolerancia a fallos por servicio',
     'al combinarse con Database per Service, la replicación protege cada dominio de '
     'forma independiente. Una falla en el almacén físico de catjard_sales no compromete '
     'a catjard_catalog ni a catjard_identity.'),
    ('RPO bajo, RTO razonable',
     'la replicación asíncrona deja una ventana mínima de pérdida (segundos en condiciones '
     'normales) y permite promoción rápida; suficiente para el SLA esperado de un proyecto '
     'académico con perfil productivo.'),
    ('Backups sin impacto en producción',
     'pg_dump y pg_basebackup pueden ejecutarse sobre la réplica sin afectar el rendimiento '
     'del nodo primario.'),
    ('Escalabilidad horizontal de lectura',
     'agregar más réplicas escala la capacidad de consulta sin rediseñar el modelo de datos.'),
]
for nombre, desc in just:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

doc.add_paragraph(
    'Se descartó la replicación lógica (publication/subscription) como mecanismo '
    'principal porque obligaría a mantener manualmente publicaciones por tabla en '
    'cada una de las seis bases de datos; la replicación física es más simple de '
    'operar y replica el cluster completo de manera transparente.'
)
doc.add_paragraph(
    'Se descartó la replicación síncrona porque introduce latencia adicional en '
    'cada commit, lo que penaliza la creación de cotizaciones y pedidos sin un '
    'beneficio proporcional para este caso de uso.'
)

# ============================================================
# 4. ADMINISTRACIÓN BD
# ============================================================
doc.add_heading('4. Administración de Base de Datos', level=1)

# 4.1 Backups
doc.add_heading('4.1 Gestión de Backups', level=2)
doc.add_paragraph(
    'La política de respaldo combina backups lógicos, backups físicos y '
    'archivado continuo de WAL para soportar Point-In-Time Recovery (PITR).'
)

backups = [
    ('Backup lógico diario',
     'pg_dump por base de datos (catjard_identity, catjard_catalog, catjard_crm, '
     'catjard_sales, catjard_inventory, catjard_operations) en formato custom (-Fc), '
     'comprimido y firmado. Programado vía Tarea Programada de Windows / cron.'),
    ('Backup físico semanal',
     'pg_basebackup completo del cluster, ejecutado sobre la réplica para no '
     'impactar al primario.'),
    ('Archivado continuo de WAL',
     'archive_mode=on; los segmentos WAL se copian a un almacenamiento aparte. '
     'Permite restaurar a cualquier punto en el tiempo (PITR) entre el último '
     'backup base y el presente.'),
    ('Retención',
     '7 días de backups diarios, 4 semanas de backups semanales, 12 meses de '
     'backups mensuales. WAL archivado se conserva el tiempo del backup base más '
     'antiguo que aún esté vigente.'),
    ('Almacenamiento',
     'copia local + copia en almacenamiento secundario (disco externo / nube). '
     'Regla 3-2-1: 3 copias, 2 medios distintos, 1 fuera del sitio.'),
    ('Pruebas de restauración',
     'simulacro mensual: se restaura un backup en una instancia de prueba y se '
     'valida que los servicios levantan correctamente y que los conteos de tablas '
     'críticas (usuarios, productos, pedidos, cotizaciones) coinciden.'),
]
for nombre, desc in backups:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

p = doc.add_paragraph()
p.add_run('Comando ejemplo (backup diario):').bold = True
cmd = """pg_dump -U postgres -h localhost -Fc -Z 9 ^
  -f "C:\\backups\\catjard_sales_2026-05-05.dump" catjard_sales

# Restaurar:
pg_restore -U postgres -h localhost -d catjard_sales_restore ^
  --clean --if-exists "C:\\backups\\catjard_sales_2026-05-05.dump" """
p = doc.add_paragraph(cmd)
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(9)

# 4.2 Seguridad
doc.add_heading('4.2 Seguridad', level=2)

doc.add_heading('A nivel de base de datos', level=3)
seg_db = [
    ('Cuentas dedicadas por servicio',
     'cada microservicio se conecta con un usuario PostgreSQL distinto '
     '(identity_app, catalog_app, crm_app, sales_app, inventory_app, '
     'operations_app), con permisos limitados a su propia base de datos '
     '(principio de mínimo privilegio).'),
    ('Sin acceso cruzado',
     'los servicios no acceden directamente a las bases de otros dominios; la '
     'comunicación inter-servicios se realiza por HTTP autenticado (Feign + JWT).'),
    ('Cifrado en tránsito',
     'conexiones JDBC con SSL (sslmode=require) en producción.'),
    ('Cifrado en reposo',
     'volúmenes con cifrado de disco a nivel de sistema operativo.'),
    ('pg_hba.conf restrictivo',
     'solo IPs/segmentos autorizados pueden conectarse, usando autenticación scram-sha-256.'),
    ('Auditoría',
     'extensión pgaudit para registrar DDL y operaciones sensibles.'),
]
for nombre, desc in seg_db:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

doc.add_heading('A nivel de aplicación', level=3)
seg_app = [
    ('Autenticación',
     'identity-service emite JWT firmado con HS256. El secreto se carga desde '
     'variable de entorno o vault, no debe vivir en el repositorio.'),
    ('Autorización',
     'cada microservicio valida el token con JwtAuthenticationFilter y aplica '
     '@PreAuthorize por rol. El frontend usa ProtectedRoute para gating por rol.'),
    ('Hash de contraseñas',
     'BCrypt (PasswordEncoder de Spring Security) en identity-service. Nunca se '
     'almacena la contraseña en claro.'),
    ('CORS controlado',
     'el API Gateway solo acepta orígenes explícitos (http://localhost:5173 en dev; '
     'el dominio público en producción).'),
    ('Headers de seguridad',
     'HSTS, X-Content-Type-Options, X-Frame-Options aplicados en el gateway.'),
    ('Validación de entrada',
     '@Valid + Bean Validation en los DTOs de cada controller para prevenir '
     'inyecciones y datos malformados.'),
    ('Protección contra SQL Injection',
     'uso exclusivo de Spring Data JPA y consultas parametrizadas; no se construyen '
     'queries por concatenación.'),
    ('Secretos fuera del código',
     'jwt.secret y credenciales de BD migrados a variables de entorno o Spring Cloud '
     'Config para producción.'),
    ('Rate limiting',
     'configurable en Spring Cloud Gateway para endpoints públicos como /api/auth/login.'),
]
for nombre, desc in seg_app:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

# 4.3 Monitoreo
doc.add_heading('4.3 Monitoreo, control y mantenimiento', level=2)

doc.add_heading('Monitoreo', level=3)
mon = [
    ('pg_stat_statements',
     'extensión activada para identificar las consultas más lentas y más frecuentes '
     'en cada base de datos.'),
    ('pg_stat_activity / pg_stat_replication',
     'monitoreo de sesiones activas, bloqueos y retraso de replicación (replication lag).'),
    ('Spring Boot Actuator',
     '/actuator/health, /actuator/metrics y /actuator/prometheus expuestos por cada '
     'microservicio para observar latencia, throughput y estado del datasource.'),
    ('Prometheus + Grafana',
     'recolección de métricas de PostgreSQL (postgres_exporter) y de los microservicios. '
     'Dashboards de conexiones, IOPS, cache hit ratio, lag de réplica y tiempos de respuesta.'),
    ('Logs centralizados',
     'logback enviando a archivo + opcionalmente ELK / Loki. Nivel DEBUG en com.catjard '
     'durante desarrollo, INFO en producción.'),
    ('Alertas',
     'umbrales en Grafana / Alertmanager para: replication_lag > 30s, conexiones > 80% del pool, '
     'errores 5xx > 1% en 5 min, espacio en disco < 20%.'),
]
for nombre, desc in mon:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

doc.add_heading('Control', level=3)
ctrl = [
    ('Versionado del esquema con Flyway',
     'cada cambio de DDL se entrega como un script V{n}__descripcion.sql en '
     'src/main/resources/db/migration. Flyway aplica los pendientes al arrancar '
     'el servicio y deja registro en flyway_schema_history.'),
    ('Política de cambios',
     'pull request obligatorio para cualquier migración; revisión de impacto sobre '
     'datos existentes; backup previo a aplicar en producción.'),
    ('Pool de conexiones',
     'HikariCP configurado por servicio con tamaño ajustado a la carga (default 10 '
     'en desarrollo). maximum-pool-size, idle-timeout y leak-detection-threshold '
     'monitoreados.'),
    ('Control de acceso',
     'roles y privilegios revisados trimestralmente; cuentas inactivas deshabilitadas.'),
]
for nombre, desc in ctrl:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

doc.add_heading('Mantenimiento', level=3)
mant = [
    ('VACUUM y ANALYZE automáticos',
     'autovacuum activado (default). Para tablas con alta tasa de UPDATE/DELETE '
     '(movimientos, tracking_eventos) se ajustan los thresholds para vacuum más '
     'agresivo.'),
    ('REINDEX periódico',
     'mensual sobre tablas grandes para evitar bloat de índices.'),
    ('Actualización de estadísticas',
     'ANALYZE manual tras cargas masivas (DataSeeder o importaciones).'),
    ('Versionado del motor',
     'PostgreSQL se mantiene en la última minor estable (parches de seguridad). '
     'Major upgrades planificados con ventana de mantenimiento y backup base previo.'),
    ('Limpieza de logs WAL',
     'archive_command + retention rotan los archivos para no llenar el disco.'),
    ('Pruebas de failover',
     'simulacros trimestrales: se promueve la réplica, se verifica que los '
     'microservicios reconectan y se reincorpora el antiguo primario como nueva réplica.'),
    ('Tareas programadas (Windows Task Scheduler)',
     'jobs nocturnos: pg_dump, validación de backups, recolección de métricas, '
     'rotación de logs.'),
]
for nombre, desc in mant:
    par = doc.add_paragraph(style='List Bullet')
    run = par.add_run(nombre)
    run.bold = True
    par.add_run(f': {desc}')

# Guardar
out = r'C:\Users\ADMIN\Documents\Universidad\datos\Cat_Jard_BD_Integracion_Replicacion_Administracion.docx'
doc.save(out)
print('OK ->', out)
